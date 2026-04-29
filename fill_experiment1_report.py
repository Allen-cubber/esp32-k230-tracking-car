from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TASK1_BLOCKS = [
    "补充实验记录：",
    "一、实验原理与命令说明。Linux 网络诊断通常围绕接口配置、路由转发和路径探测三个层面展开。`ifconfig` 或 `ip addr` 用于查看主机网卡状态、IPv4/IPv6 地址、子网掩码和链路是否启用；`ip route` 或 `route -n` 用于查看路由表和默认网关；`traceroute` 用于观察数据从源主机到目标站点所经过的逐跳转发路径，从而判断网络是否存在链路异常、路由绕行或高时延节点。",
    "二、操作步骤。1. 在实验平台进入 host1 终端，执行 `ifconfig -a` 或 `ip addr show` 查看网络设备。2. 找到 host1 对外通信使用的主接口，记录接口名称、IP 地址以及掩码/前缀长度。3. 执行 `ip route`，重点观察 `default via` 字段，对应地址即默认网关。4. 执行 `traceroute www.scut.edu.cn`，记录每一跳的 IP 地址或域名以及往返时延。5. 结合命令输出分析 host1 当前的网络连接状态和转发路径。",
    "三、任务记录与分析。",
    "1. 查询 host1 的网络设备，标记出网络接口的 IP、子网掩码。执行 `ifconfig -a` 后，可以看到 host1 的主用接口处于 `UP` 状态，接口信息中的 `inet` 字段给出 IPv4 地址，`netmask` 或 CIDR 前缀给出子网掩码。若实验平台显示为 `10.0.0.1/24`，则说明 IP 地址为 `10.0.0.1`，子网掩码为 `255.255.255.0`；若平台采用 Mininet 默认地址如 `10.0.0.1/8`，则掩码为 `255.0.0.0`。从该结果可以判断 host1 已完成地址配置，能够在本子网内进行三层通信。",
    "2. 查询 host1 的网关。执行 `ip route` 后，默认路由项通常表现为 `default via 网关地址 dev 接口名`。其中 `via` 后面的地址就是 host1 的默认网关，表示当目标地址不在本地直连网段内时，数据包将首先发往该网关，再由网关完成后续转发。默认网关的存在说明主机不仅可以访问本地局域网，还具备访问外部网络的基础条件。",
    "3. 追踪到达 `www.scut.edu.cn` 的路由路径。执行 `traceroute www.scut.edu.cn` 后，终端会依次显示从 host1 到目标服务器之间经过的多个三层设备，每一行对应一跳路由器，同时给出 3 次探测报文的时延结果。一般来说，第一跳为本地网关或实验平台出口路由器，随后经过校园网或运营商骨干节点，最终到达目标服务器所在网络。通过该命令可以直观看出报文是否能够到达目标站点、途中是否存在丢包或异常高时延节点，也能帮助分析网络拥塞和链路故障位置。",
    "四、结果总结。本任务说明 Linux 网络命令能够从不同角度反映主机的联网状态：接口命令用于确认本机配置是否正确，路由命令用于确认默认出口是否存在，路径追踪命令用于确认到目标站点的逐跳转发过程。三者结合后，可以较系统地完成网络故障排查与连接性分析。",
    "截图说明：正式提交时，建议在上述文字后分别插入 `ifconfig/ip addr`、`ip route`、`traceroute www.scut.edu.cn` 的实验截图，并保证截图中保留完整命令行输入。",
]


TASK2_BLOCKS = [
    "补充实验记录：",
    "一、实验原理。Wireshark 是常用的网络抓包与协议分析工具，可以从链路层开始逐层解析数据帧，观察以太网首部、IP 首部、传输层首部以及应用层载荷的封装关系。通过抓取 ICMP、ARP、DNS、TCP 和 HTTP 等协议报文，可以直观理解网络分层模型和典型通信过程。",
    "二、操作步骤。1. 打开 Wireshark，选择主机联网使用的接口开始抓包。2. 为避免历史数据干扰，先清空 ARP 缓存，然后分别执行 ping 子网内主机、网关和子网外主机的操作。3. 使用显示过滤器观察 `icmp`、`arp`、`tcp`、`dns`、`http` 等不同协议报文。4. 对访问网页过程进行完整抓包，定位 TCP 三次握手、HTTP 请求/响应、DNS 查询/响应以及连接关闭过程。5. 对关键报文字段进行逐项分析并记录结论。",
    "三、任务记录与分析。",
    "1. 对比 ping 子网内 IP、网关 IP 和子网外 IP 时 ICMP request/reply 帧中的源 MAC、目的 MAC。若目标主机与本机处于同一子网，则 ICMP Echo Request 的源 MAC 为本机网卡 MAC，目的 MAC 为目标主机 MAC；Echo Reply 则相反，说明同子网通信直接通过二层寻址完成。若 ping 的对象是网关，则请求帧的目的 MAC 为网关接口 MAC，回复帧的源 MAC 也为网关 MAC。若 ping 的对象位于子网外，则 IP 层的目的地址仍然是外部主机 IP，但以太网帧中的目的 MAC 不再是远端主机，而是本地默认网关的 MAC，这表明跨网段通信在链路层上先发送到下一跳，由网关继续转发。",
    "2. 分析本机与子网内邻居主机的 ARP 交互过程。当本机需要向同一子网中的另一台主机发送数据而本地 ARP 表中尚无其 MAC 地址时，会先广播 ARP Request，请求内容通常为“Who has 目标 IP? Tell 本机 IP”。该请求帧的目的 MAC 为广播地址 `ff:ff:ff:ff:ff:ff`。目标主机收到后返回 ARP Reply，告知自己的 MAC 地址，此时回复报文为单播发送给请求主机。随后，本机把获得的 IP-MAC 映射写入 ARP 缓存，后续发送 ICMP 或 TCP 数据时就可以直接封装目的 MAC，而不需要再次广播查询。",
    "3. 对步骤 8、步骤 9 过滤出的数据包进行分析。通过显示过滤器对抓包结果分类后，可以更清楚地观察协议分层特征。ARP 报文工作在链路层，不包含 IP 首部，主要字段为硬件类型、协议类型、发送端/目标端 IP 和 MAC；ICMP 报文则封装在 IP 层之上，常见类型为 Echo Request 与 Echo Reply；若过滤到 TCP/HTTP 报文，则除了以太网头和 IP 头外，还会出现 TCP 首部中的源端口、目的端口、序号、确认号、窗口大小等字段。该步骤说明不同协议承担不同网络功能，Wireshark 的过滤机制有助于针对性定位问题。",
    "4. 分析三次握手建立连接和四次挥手释放连接的包，并画出交互过程。三次握手的典型过程为：客户端先发送 `SYN` 报文，请求建立连接；服务器收到后返回 `SYN+ACK` 报文，表示同意建立连接并确认客户端序号；客户端再返回 `ACK` 报文，至此连接建立。四次挥手时，主动关闭方先发送 `FIN`；被动关闭方回复 `ACK`；待其数据发送完成后再发送 `FIN`；最后主动关闭方回复 `ACK`，连接完全断开。其本质是 TCP 的全双工连接需要分别关闭两个方向的数据流。",
    "交互过程可表示为：客户端 -> 服务器：SYN；服务器 -> 客户端：SYN, ACK；客户端 -> 服务器：ACK；客户端 -> 服务器：FIN, ACK；服务器 -> 客户端：ACK；服务器 -> 客户端：FIN, ACK；客户端 -> 服务器：ACK。",
    "5. 分析 DNS 交互过程。在访问网站之前，主机通常先向 DNS 服务器发送查询报文，请求将域名解析为 IP 地址。DNS Query 中包含事务 ID、查询名称、查询类型等字段；DNS Response 中返回对应资源记录，如 A 记录或 AAAA 记录，并在 Answers 区域中给出解析结果。若本地缓存中没有该域名映射，则必须先完成 DNS 解析，之后浏览器才会使用得到的目标 IP 发起 TCP 连接或发送 HTTP 请求。",
    "6. 分析 HTTP 请求包、应答包结构。HTTP 请求报文通常包含请求行、请求头和可选消息体。请求行中可看到方法（如 GET）、资源路径和协议版本；请求头中常见字段包括 Host、User-Agent、Accept、Connection 等。HTTP 响应报文则包含状态行、响应头和实体内容，状态行中可看到状态码（如 200 OK），响应头中常见字段包括 Content-Type、Content-Length、Server、Date 等。通过 Wireshark 展开协议层次，可以清楚看到浏览器访问网页时从 DNS、TCP 到 HTTP 的完整封装关系。",
    "四、结果总结。通过本任务可以看出，Wireshark 不仅能捕获报文，还能帮助理解二层寻址、三层转发、传输层连接管理以及应用层请求响应机制。抓包分析使抽象的协议交互过程变得可视化，是学习网络体系结构和定位故障的重要手段。",
    "截图说明：建议补充 `icmp`、`arp`、`dns`、`tcp`、`http` 过滤结果截图，以及包含三次握手和四次挥手的关键报文截图。",
]


TASK3_BLOCKS = [
    "补充实验记录：",
    "一、实验原理。IMS（IP Multimedia Subsystem，IP 多媒体子系统）是在分组交换网络上承载语音、视频和即时消息等多媒体业务的重要体系结构，其核心控制协议为 SIP。IMS 通过 CSCF 控制实体、用户数据库 HSS 以及应用服务器 AS 等模块协同工作，实现用户注册、会话控制和增值业务触发。",
    "二、实验步骤。1. 搭建 IMS 仿真环境，启动 DNS、HSS、P-CSCF、I-CSCF、S-CSCF、应用服务器 AS 以及用户终端 UE。2. 在终端侧配置用户标识、鉴权参数、归属域和代理服务器地址。3. 先完成用户注册流程抓包，再执行呼叫转移业务测试和即时消息测试。4. 在 Wireshark 中使用 `sip`、`diameter`、`dns` 等过滤器抓取并分析功能模块之间的交互消息。5. 对注册、业务触发和消息传输流程进行归纳总结。",
    "三、任务记录与分析。",
    "1. 实验系统基本架构。IMS 仿真系统通常由用户终端 UE、接入控制实体 P-CSCF、查询控制实体 I-CSCF、服务控制实体 S-CSCF、用户归属服务器 HSS 以及应用服务器 AS 组成。必要时还会包含 DNS、媒体资源功能实体和计费/日志模块。各模块分工明确：UE 负责发起注册与业务请求，P-CSCF 作为用户侧首个 SIP 接入点，I-CSCF 负责归属域查询与路由入口控制，S-CSCF 负责会话控制与业务触发，HSS 保存用户身份与业务配置，AS 负责呼叫转移、即时消息等增值业务逻辑。",
    "2. 基本架构的具体实现。从实现关系上看，UE 首先通过 P-CSCF 接入 IMS 核心网；P-CSCF 将请求转发给归属网络中的 I-CSCF；I-CSCF 查询 HSS 获得用户应绑定的 S-CSCF；S-CSCF 负责后续注册、鉴权和业务控制，并根据初始过滤条件 iFC 在需要时触发应用服务器。这样就形成了“终端接入层 + 核心控制层 + 用户数据库层 + 业务应用层”的完整 IMS 体系。该架构既便于集中管理用户，又便于灵活扩展增值业务。",
    "3. 用户登录过程。IMS 注册通常遵循 SIP REGISTER 与鉴权挑战流程。首先，UE 通过 P-CSCF 向网络发送初始 `REGISTER` 请求；I-CSCF 接收到请求后向 HSS 查询用户归属信息，并选择合适的 S-CSCF；S-CSCF 返回 `401 Unauthorized` 或 `407 Proxy Authentication Required`，要求终端携带鉴权信息重新注册；UE 根据挑战计算响应值后再次发送带 `Authorization` 头域的 `REGISTER`；网络校验通过后返回 `200 OK`，表示用户注册成功。该流程体现了 IMS 以 SIP 为核心、结合鉴权机制保证接入安全的特点。",
    "4. 呼叫转移业务交互过程。当主叫用户向被叫用户发起 `INVITE` 时，请求首先经 P-CSCF、I-CSCF 到达被叫所属的 S-CSCF。S-CSCF 根据 HSS 中保存的业务配置和 iFC 判断被叫是否开通呼叫转移业务；若业务已启用，则触发应用服务器 AS。AS 根据业务逻辑把呼叫重定向到新的目标号码或新的 SIP URI，随后网络再向新的被叫发起会话建立流程。若新被叫应答，则主叫与转移后的被叫建立通话。该过程说明 IMS 可以通过应用服务器在会话控制阶段插入业务逻辑，实现灵活的增值服务编排。",
    "5. 即时消息交互过程。IMS 中的即时消息业务通常基于 SIP SIMPLE 或 SIP MESSAGE 实现。发送方用户先完成注册，然后向接收方发送 `MESSAGE` 请求，消息体中携带文本内容；请求经 P-CSCF、S-CSCF 等控制节点转发到接收方所在侧；接收方收到消息后返回 `200 OK` 表示接收成功。若接收方不在线，则网络侧可结合应用服务器实现离线存储或延迟投递。与语音会话相比，即时消息业务不一定需要建立长期媒体会话，但仍依赖 IMS 注册状态和 SIP 路由控制完成消息分发。",
    "四、结果总结。IMS 实验表明，IMS 体系以 SIP 会话控制为基础，通过 HSS 保存用户数据，通过 CSCF 负责接入与控制，通过 AS 实现增值业务触发。用户注册、呼叫转移和即时消息三类流程共同体现了 IMS 在业务控制、身份鉴权和服务扩展方面的体系化优势。",
    "截图说明：正式提交时建议插入 IMS 拓扑图、用户注册 `REGISTER/401/REGISTER/200 OK` 抓包截图、呼叫转移 `INVITE` 触发业务截图，以及即时消息 `MESSAGE/200 OK` 交互截图。",
]


def set_run_font(run):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(12)


def format_paragraph(paragraph):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in paragraph.runs:
        set_run_font(run)


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    run = new_para.add_run(text)
    set_run_font(run)
    format_paragraph(new_para)
    return new_para


def replace_if_exact(doc, old_text, new_text):
    for para in doc.paragraphs:
        if para.text.strip() == old_text:
            para.clear()
            run = para.add_run(new_text)
            set_run_font(run)
            format_paragraph(para)
            return


def insert_block_after_text(doc, anchor_text, blocks):
    for para in doc.paragraphs:
        if para.text.strip() == anchor_text:
            current = para
            for block in blocks:
                current = insert_paragraph_after(current, block)
            return True
    return False


def main():
    if len(sys.argv) > 1:
        src = Path(sys.argv[1])
    else:
        src = Path(r"D:\HuaweiMoveData\Users\24696\Desktop\实验一.docx")
    backup = src.with_name(src.stem + "_原始备份" + src.suffix)
    doc = Document(src)

    for para in doc.paragraphs:
        format_paragraph(para)

    replace_if_exact(doc, "学生姓名：", "学生姓名：__________")
    replace_if_exact(doc, "学生学号：", "学生学号：__________")
    replace_if_exact(doc, "学生专业：", "学生专业：__________")

    inserted = [
        insert_block_after_text(doc, "3、追踪到达www.scut.edu.cn的路由路径", TASK1_BLOCKS),
        insert_block_after_text(doc, "7、步骤13，分析HTTP请求包、应答包结构。", TASK2_BLOCKS),
        insert_block_after_text(doc, "5、即时消息交互过程", TASK3_BLOCKS),
    ]

    if not all(inserted):
        raise RuntimeError("未能找到所有插入锚点，文档结构可能已变化。")

    if not backup.exists():
        backup.write_bytes(src.read_bytes())

    doc.save(src)
    print(f"updated: {src}")
    print(f"backup: {backup}")


if __name__ == "__main__":
    main()
