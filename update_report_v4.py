import os
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

docx_path = r"d:\HuaweiMoveData\Users\24696\Desktop\数据分析.docx"

try:
    doc = docx.Document(docx_path)
    print("读取现有文档...")
except Exception:
    print("未找到文档或打开失败。")
    doc = docx.Document()

# 添加新章节：结合 RapidOCR 的深度学习双轨评估法
doc.add_heading('X.5 基于预训练OCR引擎的数据集智能感知与分流策略（双轨评估体系）', level=2)

doc.add_paragraph('在利用传统计算机视觉（CV）算子（如拉普拉斯方差、HSV色域判定）完成宏观物理特征提取的基础之上，为了更加深度地契合下游模型训练的实际场景需求，本研究进一步引入了深度学习感知维度的“双轨评估体系”。本阶段摒弃了传统的底层像素统计，转而采用基于 ONNX Runtime 加速的轻量级工业推理引擎 RapidOCR，对待评估的古籍影像执行快速前向推理（Forward Inference），以此评估机器视觉网络对当前文献的“感知难度”。')

doc.add_heading('1. 感知评估维度的定义', level=3)
doc.add_paragraph('在单页抽样推理中，系统主要提取并计算两大核心深度表征参量：')
doc.add_paragraph('（1）图像级平均置信度 (Average OCR Confidence)：提取全页面所有被激活的文本边界框（Bounding Box）识别概率得分序列，并计算其数学期望。平均置信度越低，说明字迹在特征空间中发生了严重的歧义（如模糊、粘连、残缺或异体字失真）。')
p = doc.add_paragraph('（2）有效文本框召回密度 (Bounding Box Density)：基于检测网络结构（如 DBNet）在默认阈值下所成功捕获的文本边界框数量。古籍单页中极低的文本框检出数量，通常反映了该页存在大面积的物理级破损、版面缺失，或是受到严重褪色致使目标完全融入背景。')

doc.add_heading('2. 启发式自动化清洗与数据集分流池化 (Dataset Pooling)', level=3)
doc.add_paragraph('区别于随机盲抽的数据划分方式，本研究综合 RapidOCR 的感知回传数据，构建了基于极值过滤的硬性剥离规则，将超算平台上的海量扫描卷宗自动化打上了特征标签，并定向输出入两大数据集池：')

doc.add_paragraph('【挑战性验证集 / 极度破损标签】（对应文件：测试集.csv）：')
doc.add_paragraph('当系统判定当前页面的 `均值置信度 < 0.60` 或 `有效文本框数量 < 20` 时，判定其处于高度腐蚀或极度破损状态。这批数据被定向剥离输出为单独的“高优测试集”。这类“硬骨头”数据不参与日常拟合，而是专门保留用以检验开发模型在面对极端恶劣噪声与残缺场景下的抗破坏能力与鲁棒性边界（Robustness Check）。')

doc.add_paragraph('【优质基准训练集 / 完好无损标签】（对应文件：训练集.csv）：')
doc.add_paragraph('当系统判定 `均值置信度 ≥ 0.85` 且 `有效文本框数量 ≥ 50` 时，表明页面具有极高的拓扑完整性，字体密实且墨迹清晰。这部分黄金质量的数据被划拨进入主要训练池。纯净的先验数据能够引导深度学习模型在早期阶段迅速捕捉、学习古籍汉字的核心特征流形（Feature Manifold），加速网络收敛。')

doc.add_paragraph('总结而言，本报告构建了“传统CV物理特性标定”配合“深度OCR引擎逆向感知评级”的双重漏斗筛选机制。该机制成功地将 1.4TB 庞杂的无结构化大典图像集，精细解构成了具有不同训练学梯度的黄金语料库，极大拉升了整体工程的数据治理水准与科学性。')

try:
    doc.save(docx_path)
    print(f"新增的方法与原理已成功追加保存至：{docx_path}")
except Exception as e:
    print(f"写入Word出错，请检查文件是否处于打开状态 (报错信息: {e})")